import os
import json
import datetime
import inflect
import streamlit as st
import openai_completion
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Set Streamlit page configuration
st.set_page_config(
    page_title='Teaching Materials Generator',
    page_icon="logos/heig.png"
    layout='wide',
    initial_sidebar_state='collapsed'
)

# Initialize
inflect_engine = inflect.engine()
submit_button_vocab = False
submit_button_dialog = False
submit_button_tts = False

# States
if "submit_button_vocab" not in st.session_state:
    st.session_state.submit_button_vocab = False

if "submit_button_dialog" not in st.session_state:
    st.session_state.submit_button_dialog = False

if "submit_button_tts" not in st.session_state:
    st.session_state.submit_button_tts = False

if "generated_vocab" not in st.session_state:
    st.session_state.generated_vocab = ""

if "generated_dialog" not in st.session_state:
    st.session_state.generated_dialog = ""

if "words" not in st.session_state:
    st.session_state.words = ""

if "system_message_dialogs" not in st.session_state:
    st.session_state.system_message_dialogs = ""

if "generated_docx" not in st.session_state:
    st.session_state.generated_docx = ""

if "generated_tts" not in st.session_state:
    st.session_state.generated_tts = ""


def session_state_reset():
    # Reset session state variables
    st.session_state.submit_button_vocab = False
    st.session_state.submit_button_dialog = False
    st.session_state.submit_button_tts = False
    st.session_state.generated_vocab = ""
    st.session_state.generated_dialog = ""
    st.session_state.words = ""
    st.session_state.system_message_dialogs = ""
    st.session_state.generated_docx = ""
    st.session_state.generated_tts = ""


def session_state_set_vocab(bool):
    # Set submit_button_vocab session state variable
    st.session_state.submit_button_vocab = bool


def session_state_set_dialog(bool):
    # Set submit_button_dialog session state variable
    st.session_state.submit_button_dialog = bool

def session_state_set_tts(bool):
    # Set submit_button_dialog session state variable
    st.session_state.submit_button_tts = bool

def session_state_set_generated_vocab(string):
    # Set generated_vocab session state variable
    st.session_state.generated_vocab = string


def session_state_set_generated_dialog(string):
    # Set generated_dialog session state variable
    st.session_state.generated_dialog = string


def session_state_set_generated_docx(filename):
    # Set generated_docx session state variable
    st.session_state.generated_docx = filename


def session_state_set_words(string):
    # Set words session state variable
    st.session_state.words = string


def session_state_set_system_message_dialogs(string):
    # Set system_message_dialogs session state variable
    st.session_state.system_message_dialogs = string


def session_state_set_generated_tts(string):
    st.session_state.generated_tts = string


# Create a title for the app
st.title('TEACHING MATERIALS GENERATOR')
st.header('Vocabulary and dialogue for language learning', divider="red")

default_theme = 'Write the theme of the vocabulary here'

# Sidebar
st.logo('logos/heig.png')
st.sidebar.header('Settings', divider="red")

# Prompt Options
st.sidebar.subheader('Prompt Options')
taught_language = st.sidebar.selectbox('Taught Language', ['French', 'Italian', 'German'])
known_language = st.sidebar.selectbox('Known Language', ['English', 'French', 'Italian', 'German'])

# Get difficulty levels from the difficulties json file
with open('prompt/difficulty.json', 'r', encoding='utf-8') as file:
    difficulties = json.load(file)
    difficulty = st.sidebar.selectbox('Difficulty', [data['name'] for data in difficulties])
    difficulty_description = next((data['description'] for data in difficulties if data['name'] == difficulty), None)

# Number of Words, Dialogues and Turns
st.sidebar.subheader('Number of Words, Dialogues and Turns')
nb_words = st.sidebar.number_input('Number of Words', 5, 50, 10)
nb_dialogues = st.sidebar.number_input('Number of Dialogues', 1, 5, 2)
nb_turns = st.sidebar.number_input('Number of Turns', 4, 16, 6)

# Model Options
st.sidebar.subheader('Model Options')
model = st.sidebar.selectbox('Model', ['gpt-3.5-turbo', 'gpt-4o', 'gpt-4-turbo'])
temperature = st.sidebar.slider('Temperature', 0.0, 1.0, 0.8)
max_tokens = st.sidebar.slider('Max Tokens', 50, 100, 4096)

st.subheader('Vocabulary generation', divider="gray")
with open('prompt/theme.json', 'r', encoding='utf-8') as file:
    theme_data = json.load(file)
    items = [data['name'] for data in theme_data]
    theme = st.radio('Themes', items, horizontal=True, disabled=st.session_state.submit_button_vocab)
    if theme == 'custom':
        theme_description = st.text_input(default_theme)
    else:
        chosen_theme = next((data for data in theme_data if data['name'] == theme), None)
        theme_description = chosen_theme['description']
    if not st.session_state.submit_button_vocab:
        submit_button_vocab = st.button(label='Generate Vocabulary', on_click=lambda: session_state_set_vocab(True))

if st.session_state.submit_button_vocab:
    if not st.session_state.generated_vocab:
        try:
            with st.spinner('Generating...'):
                # Raise a UserWarning if theme_description is the default theme
                if theme_description == default_theme or theme_description == '':
                    raise UserWarning("No input")
                with open(f'prompt/prompt.txt', 'r', encoding='utf-8') as file:
                    template = file.read()
                system_message = template.replace('[NB_WORDS]', inflect_engine.number_to_words(nb_words))
                system_message = system_message.replace('[TAUGHT_LANGUAGE]', taught_language)
                system_message = system_message.replace('[KNOWN_LANGUAGE]', known_language)
                system_message = system_message.replace('[DIFFICULTY_DESCRIPTION]', difficulty_description)
                system_message = system_message.replace('[NB_DIALOGS]', inflect_engine.number_to_words(nb_dialogues))
                system_message = system_message.replace('[NB_TURNS]', inflect_engine.number_to_words(nb_turns))

                # Split the system_message into words and dialogs
                system_message_words, system_message_dialogs = system_message.split("-----")

                session_state_set_system_message_dialogs(system_message_dialogs)

                # Generate words using OpenAI completion
                generated_words = openai_completion.completion(system_message_words, theme_description, max_tokens,
                                                              temperature)
                # Split the generated words into lines
                words_raw = generated_words.split("\n")
                # Extract the words from the generated words
                words = " ".join(line.split("-")[0].split(" ")[2].replace(":", "") for line in words_raw)

                session_state_set_generated_vocab(generated_words)
                session_state_set_words(words)

        except UserWarning:
            # Display a warning if the theme is not provided
            st.warning("Don't forget to write the theme!")
        except Exception as e:
            # Display an error message if an error occurs
            st.error(f'An error occurred: {e}')

    # Display the generated words
    st.write('Generated Vocabulary')
    st.write(st.session_state.generated_vocab)

if st.session_state.submit_button_vocab:
    st.subheader('Dialog generation', divider="gray")
    with open('prompt/context.json', 'r', encoding='utf-8') as file:
        context_data = json.load(file)
        items = [data['name'] for data in context_data]
        context = st.radio('Context', items, horizontal=True, disabled=st.session_state.submit_button_dialog)
        if context == 'custom':
            context_description = st.text_input('Write the context of the dialogues here')
        else:
            chosen_context = next((data for data in context_data if data['name'] == context), None)
            context_description = chosen_context['description']

    with open('prompt/grammaire.json', 'r', encoding='utf-8') as file:
        grammar_data = json.load(file)
        items = [data['name'] for data in grammar_data]
        selected_items = st.multiselect('Grammar', items, disabled=st.session_state.submit_button_dialog)
        if items:
            # Add all descriptions of selected items
            grammar_description = "Use the following grammatical rules:\n\n"

            for item in selected_items:
                chosen_grammar = next((data for data in grammar_data if data['name'] == item), None)
                grammar_description += chosen_grammar['description'] + "\n\n"
        else:
            grammar_description = ""

        if not st.session_state.submit_button_dialog:
            submit_button_dialog = st.button(label='Generate Dialogs', on_click=lambda: session_state_set_dialog(True))
            reset_button = st.button(label='Reset', on_click=lambda: session_state_reset(), type="primary")
    if st.session_state.submit_button_dialog:
        if not st.session_state.generated_dialog:
            try:
                with st.spinner('Generating...'):
                    if context_description == default_theme or theme_description == '':
                        raise UserWarning("No input")
                    system_message_dialogs = st.session_state.system_message_dialogs

                    dialog_description = "The dialog is taking place in the following context: " + context_description + " using the following words:" + st.session_state.words

                    if(grammar_description != ""):
                        system_message_dialogs = system_message_dialogs.replace('[GRAMMAR]', grammar_description)
                    else:
                        system_message_dialogs = system_message_dialogs.replace('[GRAMMAR]', "")

                    generated_dialogs = openai_completion.completion(system_message_dialogs, dialog_description,
                                                                    max_tokens, temperature)

                    session_state_set_generated_dialog(generated_dialogs)

            except UserWarning:
                # Display a warning if the context is not provided
                st.warning("Don't forget to write the context!")
            except Exception as e:
                # Display an error message if an error occurs
                st.error(f'An error occurred: {e}')

        st.write('Generated Dialogs')
        st.write(st.session_state.generated_dialog)

        if not st.session_state.generated_docx:
            # save results as docx in a table
            doc = docx.Document()

            doc.add_heading(theme, level=0)

            doc.add_table(1, 6)
            #fill each tab with a random number
            table = doc.tables[0]
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Lang."
            table.cell(0, 1).text = taught_language
            table.cell(0, 2).text = difficulty
            table.cell(0, 4).text = "Date"
            table.cell(0, 5).text = datetime.datetime.now().strftime("%Y-%m-%d")

            doc.add_heading('Voc', level=1)

            doc.add_table(1, 3)
            table = doc.tables[1]
            table.cell(0, 0).text = taught_language
            table.cell(0, 1).text = known_language
            table.cell(0, 2).text = "Example"
            table.style = 'Table Grid'
            for line in st.session_state.generated_vocab.split("\n"):
                if line:
                    words = line.split(". ")[1]  # Remove the number
                    taught, known, example = words.split(" / ")
                    row_cells = table.add_row().cells
                    row_cells[0].text = taught.replace(taught_language + ": ", "")
                    row_cells[1].text = known.replace(known_language + ": ", "")
                    row_cells[2].text = example.replace("EXAMPLE: ", "")
            for cell in table.columns[2].cells:
                cell.width = docx.shared.Inches(6.0)

            doc.add_page_break()

            doc.add_heading('Dialogs', level=1)
            doc.add_paragraph("Context: " + context_description, style="Body Text 2")
            
            dialog_nb = 0

            for dialog in st.session_state.generated_dialog.split("Dialog"):
                if dialog:
                    dialog_nb += 1
                    dialog = dialog.replace(str(dialog_nb)+":\n", "")
                    doc.add_heading("Dialog " + str(dialog_nb), level=2)
                    doc.add_paragraph(dialog)
                
            doc.add_picture("logos/godeby.png",height=docx.shared.Inches(2.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save("generated_materials.docx")

            st.session_state.generated_docx = "generated_materials.docx"
            

if st.session_state.generated_docx:
                       
    st.download_button(label="Download DOCX", data=open("generated_materials.docx", 'rb').read(),
                       file_name="Godeby_CsM_" + difficulty + "_XX_tm_" + taught_language + "_" + theme + "_dialY.docx", mime="application/docx")
    
    if not st.session_state.submit_button_tts:
        submit_button_tts = st.button(label='Generate Audio', on_click=lambda: session_state_set_tts(True))
    
    if st.session_state.submit_button_tts:
        if not st.session_state.generated_tts:
            with st.spinner('Generating...'):
                generated_audio = openai_completion.tts_dialog(st.session_state.generated_dialog.replace("œ", "oe"))
                session_state_set_generated_tts(generated_audio)
        if st.session_state.generated_tts:
             st.download_button(label="Download audio", data=open("dialog.mp3", 'rb').read(),
                       file_name="dialog.mp3", mime="application/docx")
    

    reset_button = st.button(label='Reset', on_click=lambda: session_state_reset(), type="primary")
